from __future__ import annotations

import shutil
from pathlib import Path

from app.schemas.regulation import RegulationChangeDraft, RegulationParseResult


class DocxEditError(RuntimeError):
    pass


def create_revision_files(
    *,
    source_path: Path,
    output_dir: Path,
    result: RegulationParseResult,
    changes: list[RegulationChangeDraft],
) -> tuple[Path, Path, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    suffix = source_path.suffix.lower()
    document_path = output_dir / f"{source_path.stem}.ai-ready{suffix or '.bin'}"
    protocol_path = output_dir / "change_protocol.txt"
    accepted = [change for change in changes if change.status == "accepted"]
    if suffix == ".docx":
        _apply_docx(source_path, document_path, result, accepted)
        message = "Создана отредактированная DOCX-копия регламента"
    else:
        shutil.copy2(source_path, document_path)
        message = (
            "Исходный файл скопирован без изменения тела документа: адресное редактирование "
            "пока поддерживается только для DOCX"
        )
    protocol_path.write_text(_protocol(changes, message), encoding="utf-8")
    return document_path, protocol_path, message


def _apply_docx(
    source_path: Path,
    document_path: Path,
    result: RegulationParseResult,
    changes: list[RegulationChangeDraft],
) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocxEditError("Для редактирования DOCX требуется python-docx") from exc
    doc = Document(str(source_path))
    fragments = {fragment.fragmentId: fragment for fragment in result.fragments}
    for change in changes:
        fragment = fragments.get(change.targetBlockId)
        if fragment is None:
            raise DocxEditError(f"Блок {change.targetBlockId} не найден")
        location = fragment.location or {}
        paragraph_index = location.get("paragraphIndex")
        if paragraph_index is None:
            raise DocxEditError(f"Для блока {change.targetBlockId} нет DOCX-адреса абзаца")
        paragraph = doc.paragraphs[int(paragraph_index)]
        current = paragraph.text.strip()
        if fragment.contentHash and current != fragment.text.strip():
            raise DocxEditError(f"Текст блока {change.targetBlockId} изменился после распознавания")
        if change.operation in {"append_to_paragraph", "replace_text_range"}:
            _replace_paragraph_text(paragraph, change.after)
        elif change.operation in {"insert_paragraph_after", "insert_list_item"}:
            new_paragraph = _insert_paragraph_after(paragraph, change.after or change.before)
            new_paragraph.style = paragraph.style
        else:
            raise DocxEditError(f"Операция {change.operation} пока не поддерживается")
    doc.save(str(document_path))


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(text)


def _insert_paragraph_after(paragraph, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # noqa: SLF001
    new_paragraph = Paragraph(new_p, paragraph._parent)  # noqa: SLF001
    new_paragraph.add_run(text)
    return new_paragraph


def _protocol(changes: list[RegulationChangeDraft], message: str) -> str:
    lines = [message, "", "Протокол изменений", ""]
    for change in changes:
        lines.extend(
            [
                f"{change.changeId} [{change.status}]",
                f"Блок: {change.targetBlockId}",
                f"Операция: {change.operation}",
                f"Причина: {change.reason}",
                "Было:",
                change.before or "-",
                "Станет:",
                change.after or "-",
                "",
            ]
        )
    return "\n".join(lines)
