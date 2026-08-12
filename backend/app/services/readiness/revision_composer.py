from __future__ import annotations

import html
import json
import logging
import re
from pathlib import Path
from typing import Any

from app.config import settings
from app.schemas.regulation import AgentReadinessResult, RegulationParseResult, RevisionDiffBlock
from app.services.regulation.pdf_text import extract_pdf_text
from app.services.role_matching.claudehub_client import _load_json, _post_json_with_model


logger = logging.getLogger(__name__)


def create_llm_revision_files(
    *,
    source_path: Path,
    output_dir: Path,
    result: RegulationParseResult,
    readiness: AgentReadinessResult,
) -> tuple[
    Path,
    Path,
    Path | None,
    str,
    str,
    str,
    list[RevisionDiffBlock],
    list[dict],
    list[dict],
]:
    output_dir.mkdir(parents=True, exist_ok=True)
    applicable = [
        change
        for change in readiness.changes
        if change.status in {"pending", "accepted", "edited", "unchanged"}
    ]
    revised_blocks, message = _compose_revised_blocks(result, readiness, applicable)
    document_path = output_dir / f"{source_path.stem}.ai-ready.docx"
    protocol_path = output_dir / "change_protocol.txt"
    source_html = _preview_html(result.fragments, changed_ids={change.targetBlockId for change in applicable})
    revised_html = _preview_html(
        [
            fragment.model_copy(update={"text": revised_blocks.get(fragment.fragmentId, fragment.text)})
            for fragment in result.fragments
        ],
        changed_ids={change.targetBlockId for change in applicable},
    )
    diff_blocks = _diff_blocks(result, revised_blocks, applicable)
    _write_docx(document_path, result, revised_blocks, changed_ids={item.blockId for item in diff_blocks})
    pdf_path, source_preview_pages, revised_preview_pages = _write_pdf_revision_assets(
        source_path=source_path,
        output_dir=output_dir,
        result=result,
        revised_blocks=revised_blocks,
        diff_blocks=diff_blocks,
    )
    protocol_path.write_text(_protocol(readiness, diff_blocks, message), encoding="utf-8")
    return (
        document_path,
        protocol_path,
        pdf_path,
        message,
        source_html,
        revised_html,
        diff_blocks,
        source_preview_pages,
        revised_preview_pages,
    )


def _compose_revised_blocks(result: RegulationParseResult, readiness: AgentReadinessResult, changes: list) -> tuple[dict[str, str], str]:
    baseline = {fragment.fragmentId: fragment.text for fragment in result.fragments}
    fallback = _fallback_revised_blocks(baseline, changes)
    if not changes:
        return baseline, "Нет применяемых изменений; создана DOCX-копия исходного текста"
    payload = _revision_prompt(result, readiness, changes)
    try:
        raw, model = _post_json_with_model(payload, timeout=240.0)
        return _revised_blocks_from_raw(raw, baseline, fallback), _revision_model_message(model)
    except Exception as exc:  # noqa: BLE001 - keep document generation available offline.
        return _fallback_revision(baseline, fallback), (
            "ClaudeHub недоступен, применена редакция из ответов пользователя. "
            f"Ошибка моделей: {exc}"
        )


def _revision_model_message(model: str) -> str:
    if model.startswith("chad:"):
        return f"Chad API ({model.removeprefix('chad:')}) сформировал новую редакцию регламента"
    if model == settings.claudehub_external_fallback_model:
        return f"ClaudeHub GPT-5.6 Sol ({model}) сформировал новую редакцию регламента"
    return f"ClaudeHub ({model}) сформировал новую редакцию регламента"


def _revised_blocks_from_raw(raw: str, baseline: dict[str, str], fallback: dict[str, str]) -> dict[str, str]:
    data = _load_json(raw)
    revised = data.get("revisedBlocks") if isinstance(data, dict) else None
    if not isinstance(revised, list):
        raise ValueError("ClaudeHub did not return revisedBlocks")
    out = dict(baseline)
    for item in revised:
        if not isinstance(item, dict):
            continue
        block_id = str(item.get("blockId") or "")
        text = str(item.get("text") or "").strip()
        if block_id in baseline and text:
            original = baseline[block_id]
            out[block_id] = fallback.get(block_id, original) if _drops_existing_content(original, text) else text
    return out


def _drops_existing_content(original: str, revised: str) -> bool:
    original_lines = _meaningful_lines(original)
    revised_lines = _meaningful_lines(revised)
    if len(original_lines) >= 3 and len(revised_lines) < max(2, int(len(original_lines) * 0.6)):
        return True
    original_bullets = [line for line in original_lines if line.startswith(("-", "•"))]
    revised_bullets = [line for line in revised_lines if line.startswith(("-", "•"))]
    if len(original_bullets) >= 2 and len(revised_bullets) < max(1, len(original_bullets) - 1):
        return True
    if len(revised) < len(original) * 0.55:
        return True
    return False


def _meaningful_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def _fallback_revision(baseline: dict[str, str], fallback: dict[str, str]) -> dict[str, str]:
    out = dict(baseline)
    out.update(fallback)
    return out


def _revision_prompt(result: RegulationParseResult, readiness: AgentReadinessResult, changes_for_revision: list) -> dict[str, Any]:
    changed_ids = {change.targetBlockId for change in changes_for_revision}
    source_blocks = [
        {
            "blockId": fragment.fragmentId,
            "section": fragment.section,
            "text": fragment.text,
        }
        for fragment in result.fragments
        if fragment.fragmentId in changed_ids
    ]
    changes = [
        {
            "changeId": change.changeId,
            "targetBlockId": change.targetBlockId,
            "before": change.before,
            "after": change.after,
            "reason": change.reason,
            "answer": change.source.get("answer", ""),
        }
        for change in changes_for_revision
    ]
    return {
        "instruction": (
            "Сформируй новую редакцию только указанных блоков регламента. "
            "Используй исключительно подтверждённые ответы пользователя и текст исходного блока. "
            "Возвращай полный текст каждого изменяемого блока: все неизменённые пункты и подпункты должны сохраниться. "
            "Нельзя возвращать только добавленную фразу или только один подпункт многострочного блока. "
            "Не добавляй новые правила от себя, не меняй смысл незатронутых блоков. "
            "Верни только JSON с массивом revisedBlocks: [{blockId, section, text, explanation}]."
        ),
        "sourceDocument": {"fileName": result.fileName, "changedBlocks": source_blocks},
        "acceptedChanges": changes,
        "answers": [answer.model_dump(mode="json") for answer in readiness.answers],
        "responseSchema": {
            "revisedBlocks": [
                {
                    "blockId": "B-001",
                    "section": "5.2 Руководитель сектора",
                    "text": "Новая редакция пункта",
                    "explanation": "как использован ответ пользователя",
                }
            ]
        },
    }


def _write_docx(document_path: Path, result: RegulationParseResult, revised_blocks: dict[str, str], *, changed_ids: set[str]) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_COLOR_INDEX
    except ImportError as exc:
        raise RuntimeError("Для формирования DOCX требуется python-docx") from exc
    doc = Document()
    doc.add_heading(result.fileName, level=1)
    current_section = ""
    for fragment in result.fragments:
        if fragment.section and fragment.section != current_section:
            current_section = fragment.section
            doc.add_heading(current_section, level=2)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(revised_blocks.get(fragment.fragmentId, fragment.text))
        if fragment.fragmentId in changed_ids:
            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
    doc.save(str(document_path))


def _write_pdf_revision_assets(
    *,
    source_path: Path,
    output_dir: Path,
    result: RegulationParseResult,
    revised_blocks: dict[str, str],
    diff_blocks: list[RevisionDiffBlock],
) -> tuple[Path | None, list[dict], list[dict]]:
    if source_path.suffix.casefold() != ".pdf" or not source_path.exists():
        return None, [], []
    try:
        import fitz
    except ImportError:
        return None, [], []

    styled_fragments = _fragments_with_runtime_pdf_styles(source_path, result.fragments)
    changed_ids = {item.blockId for item in diff_blocks}
    fragments_by_id = {fragment.fragmentId: fragment for fragment in styled_fragments}
    changed_fragments = [
        fragment for block_id in changed_ids if (fragment := fragments_by_id.get(block_id)) is not None
    ]
    pdf_path = output_dir / f"{source_path.stem}.ai-ready.pdf"
    try:
        with fitz.open(str(source_path)) as source:
            out = fitz.open()
            replace_pages = _pages_requiring_text_replacement(source, changed_fragments, revised_blocks, result.isScan)
            for index in range(source.page_count):
                page_no = index + 1
                if page_no in replace_pages:
                    page = source[index]
                    _write_styled_page_reconstruction(
                        out,
                        page.rect,
                        [fragment for fragment in styled_fragments if fragment.page == page_no],
                        revised_blocks,
                    )
                else:
                    out.insert_pdf(source, from_page=index, to_page=index)
            for fragment in changed_fragments:
                if fragment.page in replace_pages:
                    continue
                _replace_fragment_text(out, fragment, revised_blocks.get(fragment.fragmentId, fragment.text))
            out.save(str(pdf_path), garbage=4, deflate=True)
            out.close()
    except Exception:
        logger.exception("Failed to create PDF revision assets for source=%s output=%s", source_path, pdf_path)
        return None, [], []

    source_pages = _render_pdf_preview_pages(source_path, output_dir / "preview" / "source")
    revised_pages = _render_pdf_preview_pages(pdf_path, output_dir / "preview" / "revised")
    return pdf_path, source_pages, revised_pages


def _fragments_with_runtime_pdf_styles(source_path: Path, fragments) -> list:
    if any(fragment.styleRuns for fragment in fragments):
        return list(fragments)
    try:
        extracted = extract_pdf_text(source_path)
    except Exception:
        return list(fragments)
    styled_by_page: dict[int, list] = {}
    for block in extracted.blocks:
        if block.style_runs:
            styled_by_page.setdefault(block.page, []).append(block)
    if not styled_by_page:
        return list(fragments)
    enriched = []
    for fragment in fragments:
        match = _best_style_block(fragment, styled_by_page.get(fragment.page, []))
        if match is None:
            enriched.append(fragment)
            continue
        enriched.append(
            fragment.model_copy(
                update={
                    "styleRuns": match.style_runs,
                    "bbox": list(match.bbox) if match.bbox is not None else fragment.bbox,
                    "fontSize": match.font_size or fragment.fontSize,
                    "isBold": match.is_bold or fragment.isBold,
                }
            )
        )
    return enriched


def _best_style_block(fragment, candidates) -> object | None:
    if not candidates:
        return None
    target = _style_match_text(fragment.text)
    if not target:
        return None
    scored = [(_style_match_score(target, _style_match_text(candidate.text)), candidate) for candidate in candidates]
    score, candidate = max(scored, key=lambda item: item[0])
    return candidate if score >= 0.45 else None


def _style_match_text(text: str) -> str:
    return " ".join(text.casefold().split())


def _style_match_score(left: str, right: str) -> float:
    if not left or not right:
        return 0.0
    if left in right or right in left:
        return min(len(left), len(right)) / max(len(left), len(right))
    left_tokens = set(left.split())
    right_tokens = set(right.split())
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / len(left_tokens | right_tokens)


def _pages_requiring_text_replacement(source, fragments, revised_blocks: dict[str, str], is_scan: bool) -> set[int]:
    pages: set[int] = set()
    for fragment in fragments:
        if is_scan:
            pages.add(fragment.page)
            continue
        if fragment.page < 1 or fragment.page > source.page_count:
            pages.add(fragment.page)
            continue
        page = source[fragment.page - 1]
        if not _usable_bbox(fragment.bbox, page.rect):
            pages.add(fragment.page)
            continue
        revised_text = revised_blocks.get(fragment.fragmentId, fragment.text)
        if not _text_fits_bbox(fragment, revised_text):
            pages.add(fragment.page)
    return pages


def _usable_bbox(bbox: list[float] | None, page_rect) -> bool:
    if not bbox or len(bbox) < 4:
        return False
    x0, y0, x1, y1 = [float(value) for value in bbox[:4]]
    if x1 <= x0 or y1 <= y0:
        return False
    if (x1 - x0) > page_rect.width * 0.85 and (y1 - y0) > page_rect.height * 0.85:
        return False
    return True


def _replace_fragment_text(doc, fragment, text: str) -> None:
    if fragment.page < 1 or fragment.page > doc.page_count or not _usable_bbox(fragment.bbox, doc[fragment.page - 1].rect):
        return
    import fitz

    page = doc[fragment.page - 1]
    rect = fitz.Rect(*[float(value) for value in fragment.bbox[:4]])
    page.add_redact_annot(rect, fill=(1, 1, 1))
    page.apply_redactions()
    style = _line_style(_fragment_style(fragment), text)
    fontsize = max(7.0, min(float(style["fontSize"] or fragment.fontSize or 10.0), 12.0))
    page.insert_textbox(
        rect,
        text,
        fontsize=fontsize,
        color=_pdf_color_tuple(style["color"]),
        align=0,
        **_pdf_font_kwargs(style),
    )


def _text_fits_bbox(fragment, text: str) -> bool:
    if not fragment.bbox or len(fragment.bbox) < 4:
        return False
    lines = _meaningful_lines(text)
    if len(lines) > 1 and any(_is_regulation_heading_line(line) for line in lines):
        return False
    x0, y0, x1, y1 = [float(value) for value in fragment.bbox[:4]]
    width = max(1.0, x1 - x0)
    height = max(1.0, y1 - y0)
    fontsize = max(7.0, min(float(fragment.fontSize or 10.0), 12.0))
    chars_per_line = max(12, int(width / (fontsize * 0.48)))
    required_lines = 0
    for line in lines:
        required_lines += max(1, (len(line) + chars_per_line - 1) // chars_per_line)
    available_lines = max(1, int(height / (fontsize * 1.25)))
    return required_lines <= available_lines


def _write_styled_page_reconstruction(doc, source_rect, fragments, revised_blocks: dict[str, str]) -> None:
    if not any(fragment.styleRuns for fragment in fragments):
        _write_text_pages(doc, source_rect, _page_text_from_fragments(fragments, revised_blocks))
        return
    margin = 42
    page_bottom = source_rect.height - margin
    page = doc.new_page(width=source_rect.width, height=source_rect.height)
    y_cursor = margin
    y_shift = 0.0
    flow_mode = False
    for fragment in fragments:
        text = revised_blocks.get(fragment.fragmentId, fragment.text).strip()
        if not text:
            continue
        if text == fragment.text.strip() and fragment.styleRuns:
            if flow_mode:
                page, y_cursor = _write_style_runs_in_flow(doc, page, source_rect, fragment.styleRuns, y_cursor)
            else:
                page, fragment_bottom = _write_original_style_runs(doc, page, source_rect, fragment.styleRuns, y_shift)
                y_cursor = max(y_cursor, fragment_bottom)
            continue
        was_flowing = flow_mode
        flow_mode = True
        style = _fragment_style(fragment)
        fontsize = style["fontSize"]
        line_height = fontsize * 1.35
        x = style["x"]
        y = y_cursor + fontsize if was_flowing else max(y_cursor + fontsize, style["y"] + y_shift)
        width = max(80.0, source_rect.width - x - margin)
        for paragraph in _meaningful_lines(text):
            paragraph_style = _line_style(style, paragraph)
            paragraph_fontsize = float(paragraph_style["fontSize"])
            for line in _wrap_pdf_text(paragraph, width, paragraph_fontsize):
                if y > page_bottom:
                    page = doc.new_page(width=source_rect.width, height=source_rect.height)
                    y = margin + paragraph_fontsize
                page.insert_text(
                    (x, y),
                    line,
                    fontsize=paragraph_style["fontSize"],
                    color=_pdf_color_tuple(paragraph_style["color"]),
                    **_pdf_font_kwargs(paragraph_style),
                )
                y += max(line_height, paragraph_fontsize * 1.35)
        original_bottom = _fragment_bottom(fragment) + y_shift
        if y > original_bottom:
            y_shift += y - original_bottom
        y_cursor = y + max(4.0, line_height * 0.25)


def _write_style_runs_in_flow(doc, page, source_rect, runs: list[dict], y_cursor: float) -> tuple[object, float]:
    margin = 42
    page_bottom = source_rect.height - margin
    current_page = page
    for line_runs in _group_runs_by_line(runs):
        line_style = _run_style(line_runs[0])
        line_height = line_style["fontSize"] * 1.35
        y = y_cursor + line_style["fontSize"]
        if y > page_bottom:
            current_page = doc.new_page(width=source_rect.width, height=source_rect.height)
            y = margin + line_style["fontSize"]
        for run in line_runs:
            style = _style_for_text(_run_style(run), str(run.get("text") or ""))
            current_page.insert_text(
                (style["x"], y),
                str(run.get("text") or ""),
                fontsize=style["fontSize"],
                color=_pdf_color_tuple(style["color"]),
                **_pdf_font_kwargs(style),
            )
            line_height = max(line_height, style["fontSize"] * 1.35)
        y_cursor = y + line_height * 0.15
    return current_page, y_cursor


def _group_runs_by_line(runs: list[dict]) -> list[list[dict]]:
    lines: list[list[dict]] = []
    for run in runs:
        style = _run_style(run)
        if not lines:
            lines.append([run])
            continue
        previous_y = _run_style(lines[-1][0])["y"]
        if abs(float(style["y"]) - float(previous_y)) <= 2:
            lines[-1].append(run)
        else:
            lines.append([run])
    return lines


def _write_original_style_runs(doc, page, source_rect, runs: list[dict], y_shift: float) -> tuple[object, float]:
    margin = 42
    bottom_margin = source_rect.height - margin
    max_bottom = margin
    current_page = page
    for run in runs:
        text = str(run.get("text") or "")
        if not text.strip():
            continue
        style = _run_style(run)
        x = style["x"]
        y = style["y"] + y_shift
        if y > bottom_margin:
            current_page = doc.new_page(width=source_rect.width, height=source_rect.height)
            y_shift = margin + style["fontSize"] - style["y"]
            y = style["y"] + y_shift
        current_page.insert_text(
            (x, y),
            text,
            fontsize=style["fontSize"],
            color=_pdf_color_tuple(style["color"]),
            **_pdf_font_kwargs(style),
        )
        max_bottom = max(max_bottom, _run_bottom(run) + y_shift)
    return current_page, max_bottom


def _run_style(run: dict) -> dict[str, object]:
    origin = run.get("origin") or []
    bbox = run.get("bbox") or []
    font_size = float(run.get("fontSize") or 10.0)
    style = {
        "x": float(origin[0] if len(origin) >= 2 else bbox[0] if len(bbox) >= 2 else 42.0),
        "y": float(origin[1] if len(origin) >= 2 else bbox[1] + font_size if len(bbox) >= 2 else 42.0),
        "fontSize": max(7.0, min(font_size, 18.0)),
        "fontName": str(run.get("fontName") or ""),
        "isBold": bool(run.get("isBold")),
        "isItalic": bool(run.get("isItalic")),
        "color": int(run.get("color") or 0),
    }
    return _style_for_text(style, str(run.get("text") or ""))


def _run_bottom(run: dict) -> float:
    bbox = run.get("bbox") or []
    if len(bbox) >= 4:
        return float(bbox[3])
    return float((run.get("origin") or [0, 42])[1]) + float(run.get("fontSize") or 10)


def _fragment_bottom(fragment) -> float:
    if fragment.bbox and len(fragment.bbox) >= 4:
        return float(fragment.bbox[3])
    bottoms = [_run_bottom(run) for run in fragment.styleRuns or []]
    return max(bottoms, default=42.0)


def _fragment_style(fragment) -> dict[str, object]:
    runs = fragment.styleRuns or []
    run = runs[0] if runs else {}
    bbox = fragment.bbox or run.get("bbox") or [42, 42, 0, 0]
    origin = run.get("origin") or []
    font_size = float(run.get("fontSize") or fragment.fontSize or 10.0)
    style = {
        "x": float(origin[0] if len(origin) >= 2 else bbox[0] if len(bbox) >= 2 else 42.0),
        "y": float(origin[1] if len(origin) >= 2 else bbox[1] + font_size if len(bbox) >= 2 else 42.0),
        "fontSize": max(7.0, min(font_size, 18.0)),
        "fontName": str(run.get("fontName") or ""),
        "isBold": bool(run.get("isBold") or fragment.isBold),
        "isItalic": bool(run.get("isItalic")),
        "color": int(run.get("color") or 0),
    }
    return _style_for_text(style, str(fragment.text or ""))


def _line_style(base: dict[str, object], line: str) -> dict[str, object]:
    return _style_for_text(base, line)


def _style_for_text(base: dict[str, object], text: str) -> dict[str, object]:
    style = dict(base)
    if _is_regulation_heading_line(text):
        style["isBold"] = True
        style["fontSize"] = max(float(style.get("fontSize") or 0), 12.0)
    return style


def _is_regulation_heading_line(line: str) -> bool:
    return bool(re.match(r"^\s*\d+(?:\.\d+)*\s+\S", line))


def _pdf_color_tuple(value: int) -> tuple[float, float, float]:
    red = ((value >> 16) & 255) / 255
    green = ((value >> 8) & 255) / 255
    blue = (value & 255) / 255
    return (red, green, blue)


def _write_text_pages(doc, source_rect, text: str) -> None:
    margin = 42
    fontsize = 10
    line_height = fontsize * 1.35
    width = source_rect.width - margin * 2
    bottom = source_rect.height - margin
    lines = _wrap_pdf_text(text or "Текст страницы недоступен.", width, fontsize)
    page = doc.new_page(width=source_rect.width, height=source_rect.height)
    y = margin + fontsize
    for line in lines:
        if y > bottom:
            page = doc.new_page(width=source_rect.width, height=source_rect.height)
            y = margin + fontsize
        page.insert_text((margin, y), line, fontsize=fontsize, color=(0, 0, 0), **_pdf_font_kwargs())
        y += line_height


def _wrap_pdf_text(text: str, width: float, fontsize: float) -> list[str]:
    lines: list[str] = []
    max_chars = max(24, int(width / (fontsize * 0.52)))
    for paragraph in text.splitlines():
        paragraph = paragraph.strip()
        if not paragraph:
            lines.append("")
            continue
        while len(paragraph) > max_chars:
            split_at = paragraph.rfind(" ", 0, max_chars)
            if split_at <= 0:
                split_at = max_chars
            lines.append(paragraph[:split_at].rstrip())
            paragraph = paragraph[split_at:].lstrip()
        lines.append(paragraph)
    return lines


def _pdf_font_kwargs(style: dict[str, object] | None = None) -> dict[str, str]:
    bold = bool((style or {}).get("isBold"))
    italic = bool((style or {}).get("isItalic"))
    font_name = str((style or {}).get("fontName") or "").casefold()
    prefers_times = "times" in font_name
    windows_fonts = (
        Path("C:/Windows/Fonts/timesbi.ttf") if prefers_times and bold and italic else None,
        Path("C:/Windows/Fonts/timesbd.ttf") if prefers_times and bold else None,
        Path("C:/Windows/Fonts/timesi.ttf") if prefers_times and italic else None,
        Path("C:/Windows/Fonts/times.ttf") if prefers_times else None,
        Path("C:/Windows/Fonts/arialbi.ttf") if bold and italic else None,
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else None,
        Path("C:/Windows/Fonts/ariali.ttf") if italic else None,
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf") if bold else None,
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf") if italic else None,
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    )
    for font_path in (path for path in windows_fonts if path is not None):
        if font_path.is_file():
            return {"fontname": _pdf_font_alias(font_path), "fontfile": str(font_path)}
    if bold and italic:
        return {"fontname": "hebi"}
    if bold:
        return {"fontname": "hebo"}
    if italic:
        return {"fontname": "heit"}
    return {"fontname": "helv"}


def _pdf_font_alias(font_path: Path) -> str:
    name = re.sub(r"[^a-z0-9]+", "-", font_path.stem.casefold()).strip("-")
    return f"revision-{name or 'font'}"


def _page_text_from_fragments(fragments, revised_blocks: dict[str, str]) -> str:
    lines = [
        revised_blocks.get(fragment.fragmentId, fragment.text)
        for fragment in fragments
        if (fragment.text or "").strip()
    ]
    return "\n\n".join(lines)


def _render_pdf_preview_pages(path: Path, target_dir: Path) -> list[dict]:
    try:
        import fitz
    except ImportError:
        return []
    target_dir.mkdir(parents=True, exist_ok=True)
    pages: list[dict] = []
    with fitz.open(str(path)) as doc:
        for index, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(1.4, 1.4), alpha=False)
            image_path = target_dir / f"page-{index:03d}.png"
            pix.save(str(image_path))
            pages.append({"page": index, "path": str(image_path)})
    return pages


def _preview_html(fragments, *, changed_ids: set[str]) -> str:
    items = []
    current_section = ""
    for fragment in fragments:
        if fragment.section and fragment.section != current_section:
            current_section = fragment.section
            items.append(f"<h3>{html.escape(current_section)}</h3>")
        cls = "changed" if fragment.fragmentId in changed_ids else "plain"
        items.append(f'<div class="block {cls}">{html.escape(fragment.text)}</div>')
    return "\n".join(
        [
            "<style>",
            "body{font-family:Segoe UI,Arial,sans-serif;color:#101817;background:#fff;}",
            ".block{padding:10px 12px;margin:8px 0;border-radius:12px;border:1px solid rgba(16,24,23,.08);}",
            ".changed{background:rgba(8,116,95,.08);border-color:rgba(8,116,95,.28);}",
            "h3{font-size:14px;margin:18px 0 8px;color:#06483D;}",
            "</style>",
            *items,
        ]
    )


def _diff_blocks(result: RegulationParseResult, revised_blocks: dict[str, str], accepted: list) -> list[RevisionDiffBlock]:
    by_id = {fragment.fragmentId: fragment for fragment in result.fragments}
    changed_ids = {change.targetBlockId for change in accepted}
    diff = []
    for block_id in changed_ids:
        fragment = by_id.get(block_id)
        before = fragment.text if fragment is not None else ""
        after = revised_blocks.get(block_id, before)
        if before.strip() == after.strip():
            continue
        diff.append(
            RevisionDiffBlock(
                blockId=block_id,
                section=fragment.section if fragment is not None else "",
                before=before,
                after=after,
                page=fragment.page if fragment is not None else 0,
                bbox=fragment.bbox if fragment is not None else None,
                status="changed",
            )
        )
    return diff


def _fallback_revised_blocks(baseline: dict[str, str], changes: list) -> dict[str, str]:
    out = dict(baseline)
    for change in changes:
        block_id = change.targetBlockId
        if not block_id or not change.after.strip():
            continue
        before = (change.before or baseline.get(block_id, "")).strip()
        after = change.after.strip()
        addition = _addition_from_after(before, after)
        if not addition:
            out[block_id] = after
            continue
        current = out.get(block_id, before).strip()
        if addition not in current:
            out[block_id] = f"{current.rstrip()} {addition}".strip()
    return out


def _addition_from_after(before: str, after: str) -> str:
    if before and after.startswith(before):
        return after[len(before) :].strip()
    return after.strip()


def _protocol(readiness: AgentReadinessResult, diff_blocks: list[RevisionDiffBlock], message: str) -> str:
    payload = {
        "message": message,
        "readinessRunId": readiness.readinessRunId,
        "answers": [answer.model_dump(mode="json") for answer in readiness.answers],
        "diffBlocks": [item.model_dump(mode="json") for item in diff_blocks],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)
