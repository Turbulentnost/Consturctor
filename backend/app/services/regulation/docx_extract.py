from __future__ import annotations

import hashlib
from pathlib import Path

from app.services.regulation.types import ExtractedBlock, ExtractedDocument, ExtractedTable


def extract_docx(path: Path) -> ExtractedDocument:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("Для DOCX требуется зависимость python-docx") from exc

    doc = Document(str(path))
    blocks: list[ExtractedBlock] = []
    current_section = ""
    page = 1
    paragraph_indexes = {
        paragraph._element: index  # noqa: SLF001 - python-docx addresses paragraphs by XML node
        for index, paragraph in enumerate(doc.paragraphs)
    }
    table_indexes = {
        table._element: index  # noqa: SLF001 - python-docx addresses tables by XML node
        for index, table in enumerate(doc.tables)
    }

    for item in _iter_body_items(doc, Paragraph, Table, qn):
        if isinstance(item, Paragraph):
            paragraph_index = paragraph_indexes.get(item._element, 0)  # noqa: SLF001
            text = item.text.strip()
            if not text:
                if _paragraph_has_page_break(item):
                    page += 1
                continue
            style_name = item.style.name or "" if item.style is not None else ""
            style = style_name.lower()
            list_level = _numbering_level(item)
            is_list = "list" in style or "список" in style or list_level >= 0
            kind = "list" if is_list else "text"
            is_heading = "heading" in style or "заголовок" in style
            style_runs = _paragraph_style_runs(item)
            font_size = _style_runs_font_size(style_runs) or _paragraph_font_size(item)
            is_bold = _style_runs_is_bold(style_runs) or is_heading
            numbering = _leading_numbering(text)
            layout = _paragraph_layout(item, list_level=list_level)
            if is_heading:
                current_section = text
            blocks.append(
                ExtractedBlock(
                    page=page,
                    section=current_section,
                    text=text,
                    kind=kind,
                    block_type="heading" if is_heading else ("list_item" if kind == "list" else "paragraph"),
                    font_size=font_size,
                    is_bold=is_bold,
                    numbering=numbering,
                    confidence=1.0,
                    location={
                        "documentPart": "word/document.xml",
                        "paragraphIndex": paragraph_index,
                        "paragraphXmlId": _paragraph_xml_id(item),
                        "sectionPath": [current_section] if current_section else [],
                        **layout,
                    },
                    style=style_name,
                    style_runs=style_runs,
                    content_hash=_content_hash(text),
                )
            )
            if _paragraph_has_page_break(item):
                page += 1
            continue

        table_index = table_indexes.get(item._element, 0)  # noqa: SLF001
        rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
        if not rows:
            continue
        headers = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        text = "\n".join(" | ".join(cell for cell in row if cell) for row in rows)
        blocks.append(
            ExtractedBlock(
                page=max(1, page),
                section=current_section,
                text=text,
                kind="table",
                block_type="table",
                table=ExtractedTable(headers=headers, rows=body),
                confidence=1.0,
                location={"documentPart": "word/document.xml", "tableIndex": table_index},
                style="table",
                content_hash=_content_hash(text),
            )
        )

    return ExtractedDocument(page_count=max(1, page), blocks=blocks)


def _iter_body_items(doc, paragraph_cls, table_cls, qn):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield paragraph_cls(child, doc)
        elif child.tag == qn("w:tbl"):
            yield table_cls(child, doc)


def _paragraph_style_runs(paragraph) -> list[dict[str, object]]:
    runs: list[dict[str, object]] = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        font_name = _run_font_name(run)
        font_size = _run_font_size(run, paragraph)
        runs.append(
            {
                "text": text,
                "fontName": font_name,
                "fontSize": float(font_size or 0),
                "isBold": _run_is_bold(run, paragraph, font_name),
                "isItalic": bool(run.italic) or _font_marker_has(font_name, ("italic", "oblique")),
                "underline": bool(run.underline),
                "color": _run_color(run),
            }
        )
    return runs


def _paragraph_layout(paragraph, *, list_level: int) -> dict[str, object]:
    left = _effective_length_pt(paragraph, "left_indent")
    first_line = _effective_length_pt(paragraph, "first_line_indent")
    right = _effective_length_pt(paragraph, "right_indent")
    space_before = _effective_length_pt(paragraph, "space_before")
    space_after = _effective_length_pt(paragraph, "space_after")
    if left <= 0 and first_line == 0 and list_level >= 0:
        left = 18.0 * (list_level + 1)
    return {
        "indentPt": left,
        "firstLineIndentPt": first_line,
        "rightIndentPt": right,
        "spaceBeforePt": space_before,
        "spaceAfterPt": space_after,
        "alignment": _alignment_name(paragraph),
        "listLevel": list_level,
    }


def _effective_length_pt(paragraph, attr: str) -> float:
    direct = getattr(paragraph.paragraph_format, attr, None)
    if direct is not None:
        return _length_pt(direct)
    style = paragraph.style
    if style is None:
        return 0.0
    return _length_pt(getattr(style.paragraph_format, attr, None))


def _length_pt(value) -> float:
    if value is None:
        return 0.0
    try:
        return float(value.pt)
    except Exception:
        return 0.0


def _alignment_name(paragraph) -> str:
    alignment = paragraph.paragraph_format.alignment
    if alignment is None and paragraph.style is not None:
        alignment = paragraph.style.paragraph_format.alignment
    if alignment is None:
        return ""
    name = getattr(alignment, "name", None)
    if name:
        return str(name).lower()
    return str(alignment).split("(")[0].strip().lower().rsplit(".", 1)[-1]


def _numbering_level(paragraph) -> int:
    try:
        properties = paragraph._element.pPr  # noqa: SLF001
        if properties is None or properties.numPr is None:
            return -1
        level = properties.numPr.ilvl
        if level is None or level.val is None:
            return 0
        return int(level.val)
    except Exception:
        return -1


def _run_font_name(run) -> str:
    name = run.font.name if run.font is not None else None
    if name:
        return str(name)
    try:
        from docx.oxml.ns import qn

        properties = run._element.rPr  # noqa: SLF001
        if properties is None:
            return ""
        fonts = properties.rFonts
        if fonts is None:
            return ""
        for key in (qn("w:eastAsia"), qn("w:cs"), qn("w:hAnsi"), qn("w:ascii")):
            value = fonts.get(key)
            if value:
                return str(value)
    except Exception:
        return ""
    return ""


def _run_font_size(run, paragraph) -> float | None:
    if run.font is not None and run.font.size is not None:
        return float(run.font.size.pt)
    return _paragraph_font_size(paragraph)


def _paragraph_font_size(paragraph) -> float | None:
    style = paragraph.style
    if style is not None and style.font is not None and style.font.size is not None:
        return float(style.font.size.pt)
    return None


def _run_is_bold(run, paragraph, font_name: str) -> bool:
    if run.bold is True:
        return True
    if _font_marker_has(font_name, ("bold", "black", "demi", "semibold", "-bd")):
        return True
    if run.bold is False:
        return False
    style = paragraph.style
    if style is not None and style.font is not None and style.font.bold:
        return True
    return False


def _run_color(run) -> int:
    try:
        rgb = run.font.color.rgb if run.font is not None and run.font.color is not None else None
        if rgb is None:
            return 0
        return int(str(rgb), 16)
    except Exception:
        return 0


def _font_marker_has(font_name: str, markers: tuple[str, ...]) -> bool:
    marker = font_name.casefold()
    return any(item in marker for item in markers)


def _style_runs_font_size(style_runs: list[dict[str, object]]) -> float | None:
    sizes = [float(run.get("fontSize") or 0) for run in style_runs if float(run.get("fontSize") or 0) > 0]
    return max(sizes) if sizes else None


def _style_runs_is_bold(style_runs: list[dict[str, object]]) -> bool:
    return bool(style_runs) and any(bool(run.get("isBold")) for run in style_runs)


def _paragraph_has_page_break(paragraph) -> bool:
    for run in paragraph.runs:
        xml = run._element.xml  # noqa: SLF001 - python-docx exposes break details only via XML.
        if 'w:type="page"' in xml or "lastRenderedPageBreak" in xml:
            return True
    return False


def _paragraph_xml_id(paragraph) -> str:
    element = paragraph._element  # noqa: SLF001 - python-docx exposes XML ids only here.
    for key in ("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", "w14:paraId"):
        value = element.get(key)
        if value:
            return str(value)
    return ""


def _content_hash(text: str) -> str:
    return "sha256:" + hashlib.sha256((text or "").encode("utf-8")).hexdigest()


def _leading_numbering(text: str) -> str | None:
    parts = text.strip().split(maxsplit=1)
    if not parts:
        return None
    first = parts[0].rstrip(".)")
    if first and all(piece.isdigit() for piece in first.split(".")):
        return first
    return None
