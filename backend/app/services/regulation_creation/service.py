from __future__ import annotations

import ast
import json
import re
import tempfile
from pathlib import Path
from collections.abc import Iterator
from uuid import uuid4

from sqlalchemy.orm import Session

from app.config import settings
from app.models.regulation import (
    RegulationCreationDraft,
    RegulationCreationMessage,
    RegulationDocument,
)
from app.models.user import AppUser
from app.schemas.regulation import (
    RegulationCreationApplyRequest,
    RegulationCreationMessage as CreationMessageSchema,
    RegulationCreationSendRequest,
    RegulationCreationSession,
    RegulationCreationTurn,
    RegulationFragment,
    RegulationParseResult,
)
from app.services.regulation import RegulationError
from app.services.regulation.storage import new_regulation_id, save_upload
from app.services.regulation_creation.cursor_agent import (
    CursorAgentError,
    archive_agent,
    cancel_run,
    create_agent,
    create_run,
    stream_run_events,
    wait_for_run,
)
from app.services.regulation.full_text import compose_regulation_text
from app.services.regulation.detect import is_scan_pdf
from app.services.regulation.pdf_ocr import extract_pdf_scan
from app.services.regulation_creation.interview import (
    append_user_turn,
    build_creation_prompt,
    build_followup_creation_prompt,
    creation_system_rules,
    document_from_interview,
    document_has_body,
    document_has_full_text,
    interview_sdk_agent_id,
    interview_snapshot,
    is_replacement_garbage,
    merge_agent_payload,
    new_interview_state,
    ready_blocker,
    set_interview_position,
    set_sdk_agent_id,
)
from app.services.workflows.document import DocumentError, load_attachment_bytes

_CREATION_ATTACH_SUFFIXES = {".doc", ".docx", ".pdf", ".md", ".txt"}
_MAX_ATTACH_CHARS = 120_000
_MESSAGE_CONTENT_LIMIT = 7900


FIRST_QUESTION = (
    "Приложите один или несколько файлов с обязанностями/процессами или коротко напишите должность "
    "и функции пользователя. Я разберу документы и буду уточнять каждый пробел по одному."
)
class RegulationCreationError(Exception):
    def __init__(self, message: str, status_code: int = 400) -> None:
        super().__init__(message)
        self.message = message
        self.status_code = status_code


OPEN_CREATION_STATUSES = ("collecting_positions", "interview", "generating", "error")


def get_active_creation_session(db: Session, *, user_id: str) -> RegulationCreationSession | None:
    draft = (
        db.query(RegulationCreationDraft)
        .filter(
            RegulationCreationDraft.user_id == user_id,
            RegulationCreationDraft.status.in_(list(OPEN_CREATION_STATUSES)),
        )
        .order_by(RegulationCreationDraft.updated_at.desc())
        .first()
    )
    if draft is None:
        return None
    return _session(db, draft)


def start_creation_session(
    db: Session,
    *,
    user_id: str,
    fresh: bool = False,
) -> RegulationCreationSession:
    if not fresh:
        existing = get_active_creation_session(db, user_id=user_id)
        if existing is not None:
            return existing
    else:
        terminate_active_creation_sessions(db, user_id=user_id)

    user = db.get(AppUser, user_id)
    interview = set_interview_position(new_interview_state(), getattr(user, "position", "") or "")
    draft = RegulationCreationDraft(
        id=f"reg-create-{uuid4().hex[:12]}",
        user_id=user_id,
        status="collecting_positions",
        style_profile_json={},
        interview_json=interview,
    )
    db.add(draft)
    db.flush()
    _add_message(db, draft=draft, role="assistant", content=FIRST_QUESTION)
    db.commit()
    db.refresh(draft)
    return _session(db, draft)


def get_creation_session(db: Session, *, user_id: str, draft_id: str) -> RegulationCreationSession:
    return _session(db, _get_draft(db, user_id=user_id, draft_id=draft_id))


def peek_creation_turn(db: Session, *, user_id: str, draft_id: str) -> RegulationCreationTurn:
    draft = _get_draft(db, user_id=user_id, draft_id=draft_id)
    last_user = ""
    for item in reversed(_messages_for_draft(db, draft.id)):
        if item.role == "user":
            last_user = item.content or ""
            break
    return _turn_payload(
        db,
        draft,
        message=last_user,
        force_create=_is_force_create_message(last_user),
    )


def get_creation_document(db: Session, *, user_id: str, draft_id: str) -> Path:
    draft = _get_draft(db, user_id=user_id, draft_id=draft_id)
    path = Path(draft.result_document_path or "")
    if path.is_file():
        return path
    if draft.result_regulation_id:
        from app.services.regulation.storage import get_document

        doc = get_document(db, regulation_id=draft.result_regulation_id, user_id=user_id)
        stored = Path((doc.storage_path if doc is not None else "") or "")
        if stored.is_file():
            return stored
    rebuilt = _rebuild_creation_docx(draft)
    if rebuilt is not None and rebuilt.is_file():
        draft.result_document_path = str(rebuilt)
        db.add(draft)
        db.commit()
        return rebuilt
    raise RegulationCreationError("Файл регламента ещё не создан", status_code=404)


def _rebuild_creation_docx(draft: RegulationCreationDraft) -> Path | None:
    document = draft.draft_document_json if isinstance(draft.draft_document_json, dict) else {}
    if not document_has_body(document):
        title = str((document or {}).get("title") or "").strip()
        document = document_from_interview(draft.interview_json, title)
    if not document_has_body(document):
        return None
    output_dir = settings.regulation_storage_dir / "created" / draft.id
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / _safe_filename(str(document.get("title") or "created-regulation"))
    path = path.with_suffix(".docx")
    _write_docx(path, document)
    return path


def persist_creation_turn(
    db: Session,
    *,
    user_id: str,
    draft_id: str,
    request: RegulationCreationSendRequest,
    files: list[tuple[str, bytes]] | None = None,
) -> RegulationCreationTurn:
    attachments = _load_creation_attachments(files or [])
    message = request.message.strip()
    if not message and not attachments:
        raise RegulationCreationError("Введите сообщение или приложите файл")
    draft = _get_draft(db, user_id=user_id, draft_id=draft_id)
    if draft.status == "finalized":
        return _turn_payload(db, draft, message="", force_create=False)
    force_create = _is_force_create_message(message)
    display_message = _display_user_message(message, attachments)
    draft.interview_json = append_user_turn(draft.interview_json, message, attachments)
    _add_message(
        db,
        draft=draft,
        role="user",
        content=display_message,
        structured=_attachments_structured(attachments),
    )
    draft.status = "generating"
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return _turn_payload(db, draft, message=message, force_create=force_create)


def apply_creation_reply(
    db: Session,
    *,
    user_id: str,
    draft_id: str,
    request: RegulationCreationApplyRequest,
) -> RegulationCreationSession:
    draft = _get_draft(db, user_id=user_id, draft_id=draft_id)
    if draft.status == "finalized":
        return _session(db, draft)
    if request.sdkAgentId.strip():
        draft.interview_json = set_sdk_agent_id(draft.interview_json, request.sdkAgentId)
    force_create = bool(request.forceCreate)
    _apply_agent_reply(
        db,
        user_id=user_id,
        draft=draft,
        raw=request.answer,
        force_create=force_create,
    )
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return _session(db, draft)


def _turn_payload(
    db: Session,
    draft: RegulationCreationDraft,
    *,
    message: str,
    force_create: bool,
) -> RegulationCreationTurn:
    sdk_id = interview_sdk_agent_id(draft.interview_json)
    prompt = (
        build_followup_creation_prompt(message=message, force_create=force_create)
        if sdk_id
        else build_creation_prompt(
            state=draft.interview_json,
            message=message,
            initial=True,
            force_create=force_create,
            include_attachment_bodies=False,
        )
    )
    return RegulationCreationTurn(
        session=_session(db, draft),
        interview=interview_snapshot(draft.interview_json),
        sdkPrompt=prompt,
        sdkRules=creation_system_rules(force_create=force_create),
        sdkAgentId=sdk_id,
        forceCreate=force_create,
    )


def terminate_active_creation_sessions(db: Session, *, user_id: str) -> dict:
    drafts = (
        db.query(RegulationCreationDraft)
        .filter(
            RegulationCreationDraft.user_id == user_id,
            RegulationCreationDraft.status.in_(list(OPEN_CREATION_STATUSES)),
        )
        .all()
    )
    closed = 0
    errors: list[str] = []
    for draft in drafts:
        if draft.cursor_agent_id:
            if draft.latest_run_id:
                try:
                    cancel_run(draft.cursor_agent_id, draft.latest_run_id)
                except CursorAgentError as exc:
                    if exc.status_code != 409:
                        errors.append(exc.message)
            try:
                archive_agent(draft.cursor_agent_id)
            except CursorAgentError as exc:
                errors.append(exc.message)
        draft.status = "closed"
        db.add(draft)
        closed += 1
    db.commit()
    return {"closed": closed, "errors": errors}


def send_creation_message(
    db: Session,
    *,
    user_id: str,
    draft_id: str,
    request: RegulationCreationSendRequest,
    files: list[tuple[str, bytes]] | None = None,
) -> RegulationCreationSession:
    attachments = _load_creation_attachments(files or [])
    message = request.message.strip()
    if not message and not attachments:
        raise RegulationCreationError("Введите сообщение или приложите файл")
    draft = _get_draft(db, user_id=user_id, draft_id=draft_id)
    if draft.status == "finalized":
        return _session(db, draft)
    force_create = _is_force_create_message(message)
    display_message = _display_user_message(message, attachments)
    draft.interview_json = append_user_turn(draft.interview_json, message, attachments)
    _add_message(
        db,
        draft=draft,
        role="user",
        content=display_message,
        structured=_attachments_structured(attachments),
    )
    draft.status = "generating"
    db.add(draft)
    db.commit()

    prompt = build_creation_prompt(
        state=draft.interview_json,
        message=message,
        initial=not draft.cursor_agent_id,
        force_create=force_create,
    )
    try:
        if not draft.cursor_agent_id:
            agent_id, run_id = create_agent(prompt)
            draft.cursor_agent_id = agent_id
            draft.latest_run_id = run_id
        else:
            run_id = create_run(draft.cursor_agent_id, prompt)
            draft.latest_run_id = run_id
        db.add(draft)
        db.commit()
        run = wait_for_run(draft.cursor_agent_id, draft.latest_run_id)
    except CursorAgentError as exc:
        draft.status = "error"
        db.add(draft)
        db.commit()
        raise RegulationCreationError(exc.message, status_code=exc.status_code) from exc

    _apply_agent_reply(
        db,
        user_id=user_id,
        draft=draft,
        raw=str(run.get("result") or ""),
        force_create=force_create,
    )
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return _session(db, draft)


def stream_creation_message(
    db: Session,
    *,
    user_id: str,
    draft_id: str,
    request: RegulationCreationSendRequest,
    files: list[tuple[str, bytes]] | None = None,
) -> Iterator[dict]:
    attachments = _load_creation_attachments(files or [])
    message = request.message.strip()
    if not message and not attachments:
        raise RegulationCreationError("Введите сообщение или приложите файл")
    draft = _get_draft(db, user_id=user_id, draft_id=draft_id)
    if draft.status == "finalized":
        yield {"type": "session", "session": _session(db, draft).model_dump(mode="json")}
        return

    force_create = _is_force_create_message(message)
    display_message = _display_user_message(message, attachments)
    draft.interview_json = append_user_turn(draft.interview_json, message, attachments)
    _add_message(
        db,
        draft=draft,
        role="user",
        content=display_message,
        structured=_attachments_structured(attachments),
    )
    draft.status = "generating"
    db.add(draft)
    db.commit()
    yield {"type": "status", "status": "generating"}

    prompt = build_creation_prompt(
        state=draft.interview_json,
        message=message,
        initial=not draft.cursor_agent_id,
        force_create=force_create,
    )
    final_text = ""
    assistant_parts: list[str] = []
    try:
        if not draft.cursor_agent_id:
            agent_id, run_id = create_agent(prompt)
            draft.cursor_agent_id = agent_id
            draft.latest_run_id = run_id
        else:
            run_id = create_run(draft.cursor_agent_id, prompt)
            draft.latest_run_id = run_id
        db.add(draft)
        db.commit()
        try:
            for event in stream_run_events(draft.cursor_agent_id, draft.latest_run_id):
                event_type = str(event.get("event") or "")
                data = event.get("data") if isinstance(event.get("data"), dict) else {}
                if event_type == "thinking":
                    text = str(data.get("text") or "")
                    if text:
                        yield {"type": "thinking", "text": text}
                elif event_type == "assistant":
                    text = str(data.get("text") or "")
                    if text:
                        assistant_parts.append(text)
                        yield {"type": "assistant", "text": text}
                elif event_type == "result":
                    final_text = str(data.get("text") or "")
                    status = str(data.get("status") or "")
                    if status and status != "FINISHED":
                        raise CursorAgentError(f"Cursor Agent завершился со статусом {status}", status_code=502)
        except CursorAgentError as exc:
            if exc.status_code != 409 and "stream_unavailable" not in exc.message:
                raise
            yield {"type": "status", "status": "stream_unavailable_polling"}
    except CursorAgentError as exc:
        draft.status = "error"
        db.add(draft)
        db.commit()
        yield {"type": "error", "message": exc.message}
        return

    if not final_text:
        final_text = "".join(assistant_parts).strip()
    if not final_text:
        try:
            final_text = str(wait_for_run(draft.cursor_agent_id, draft.latest_run_id).get("result") or "")
        except CursorAgentError as exc:
            draft.status = "error"
            db.add(draft)
            db.commit()
            yield {"type": "error", "message": exc.message}
            return

    _apply_agent_reply(db, user_id=user_id, draft=draft, raw=final_text, force_create=force_create)
    db.commit()
    db.refresh(draft)
    yield {"type": "session", "session": _session(db, draft).model_dump(mode="json")}


def _apply_agent_reply(
    db: Session,
    *,
    user_id: str,
    draft: RegulationCreationDraft,
    raw: str,
    force_create: bool = False,
) -> None:
    raw = raw.strip()
    parsed = _parse_agent_response(raw)
    if is_replacement_garbage(parsed.get("message")) or is_replacement_garbage(raw):
        parsed["message"] = (
            "Ответ агента пришёл в нечитаемом виде. Нажмите отправку ещё раз "
            "или коротко напишите должность и первую функцию."
        )
        parsed["quickAnswers"] = [
            "Повторить разбор файлов",
            "Опишу функции сообщением",
        ]
        parsed["status"] = "need_more"
    draft.interview_json = merge_agent_payload(draft.interview_json, parsed)
    blocker = None if force_create else ready_blocker(parsed, draft.interview_json)
    if blocker is not None:
        _add_message(
            db,
            draft=draft,
            role="assistant",
            content=blocker.message,
            structured={
                "quickAnswers": blocker.quick_answers,
                "blockedReady": {
                    "functionId": blocker.function_id,
                    "field": blocker.field,
                },
            },
        )
        draft.status = "interview"
        if positions := parsed.get("positions"):
            draft.positions_json = [str(item) for item in positions if str(item).strip()]
        db.add(draft)
        return
    document = parsed.get("document") if isinstance(parsed.get("document"), dict) else None
    wants_document = parsed.get("status") == "ready" or force_create
    has_full_document = document_has_full_text(document)
    if wants_document and not has_full_document and not force_create:
        state = draft.interview_json if isinstance(draft.interview_json, dict) else {}
        draft.interview_json = {**state, "document_write_required": True}
        draft.status = "interview"
        if positions := parsed.get("positions"):
            draft.positions_json = [str(item) for item in positions if str(item).strip()]
        db.add(draft)
        return
    if wants_document and force_create and not document_has_body(document):
        title = str((document or {}).get("title") or "").strip()
        document = document_from_interview(draft.interview_json, title)
    if wants_document and force_create and not document_has_body(document):
        document = _stub_document(parsed, title=str((document or {}).get("title") or "").strip())
    if wants_document and (has_full_document or (force_create and document_has_body(document))):
        if isinstance(draft.interview_json, dict):
            draft.interview_json.pop("document_write_required", None)
        try:
            result = _finalize_document(db, user_id=user_id, draft=draft, document=document or {})
        except RegulationError as exc:
            draft.status = "error"
            db.add(draft)
            db.commit()
            raise RegulationCreationError(exc.message, status_code=exc.status_code) from exc
        except RegulationCreationError:
            draft.status = "error"
            db.add(draft)
            db.commit()
            raise
        _add_message(
            db,
            draft=draft,
            role="assistant",
            content=parsed.get("message") or "Регламент сформирован. Проверьте документ перед созданием агента.",
            structured={"resultRegulationId": result.regulationId, "document": document},
        )
        draft.status = "finalized"
        draft.result_regulation_id = result.regulationId
    else:
        quick_answers = _quick_answers(parsed.get("quickAnswers"))
        _add_message(
            db,
            draft=draft,
            role="assistant",
            content=parsed.get("message") or raw or "Уточните, пожалуйста, детали процесса.",
            structured={"quickAnswers": quick_answers},
        )
        draft.status = "interview"
        if positions := parsed.get("positions"):
            draft.positions_json = [str(item) for item in positions if str(item).strip()]


def _finalize_document(
    db: Session,
    *,
    user_id: str,
    draft: RegulationCreationDraft,
    document: dict,
) -> RegulationParseResult:
    output_dir = settings.regulation_storage_dir / "created" / draft.id
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / _safe_filename(str(document.get("title") or "created-regulation"))
    path = path.with_suffix(".docx")
    _write_docx(path, document)
    regulation_id = new_regulation_id()
    stored = save_upload(regulation_id=regulation_id, filename=path.name, data=path.read_bytes())
    result = _result_from_created_document(
        regulation_id=regulation_id,
        filename=path.name,
        document=document,
    )
    db.add(
        RegulationDocument(
            id=regulation_id,
            user_id=user_id,
            file_name=path.name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path=str(stored),
            is_scan=False,
            result_json=result.model_dump(mode="json"),
        )
    )
    draft.result_document_path = str(path)
    draft.draft_document_json = document
    return result


def _result_from_created_document(
    *,
    regulation_id: str,
    filename: str,
    document: dict,
) -> RegulationParseResult:
    title = str(document.get("title") or "Регламент").strip() or "Регламент"
    fragments: list[RegulationFragment] = []
    sections: list[str] = []
    index = 0
    for section, section_title, section_path, _level in _iter_document_sections(
        document.get("sections") or [],
        parent_path=[title],
    ):
        section_title = section_title or f"Раздел {index + 1}"
        if section_title not in sections:
            sections.append(section_title)
        texts = [section_title] if section_title else []
        for paragraph in section.get("paragraphs") or []:
            value = str(paragraph or "").strip()
            if value:
                texts.append(value)
        for item in section.get("items") or []:
            value = _document_item_text(item)
            if value:
                texts.append(value)
        for text in texts:
            index += 1
            fragments.append(
                RegulationFragment(
                    fragmentId=f"{regulation_id}-block-{index:02d}",
                    page=1,
                    section=section_title,
                    sectionPath=section_path or [title, section_title],
                    kind="text",
                    blockType="heading" if text == section_title else "paragraph",
                    text=text,
                    isBold=text == section_title,
                )
            )
    if not fragments:
        fragments.append(
            RegulationFragment(
                fragmentId=f"{regulation_id}-block-01",
                page=1,
                section=title,
                sectionPath=[title],
                kind="text",
                blockType="heading",
                text=title,
                isBold=True,
            )
        )
        sections = [title]
    return RegulationParseResult(
        regulationId=regulation_id,
        fileName=filename,
        pageCount=1,
        sectionCount=len(sections),
        recognitionQuality=1.0,
        isScan=False,
        sections=sections,
        fragments=fragments,
    )


def _stub_document(parsed: dict, *, title: str) -> dict:
    message = str(parsed.get("message") or "").strip() or "Регламент сформирован по текущим данным."
    return {
        "title": title or "Регламент",
        "sections": [
            {
                "number": "1",
                "title": "Порядок",
                "paragraphs": [message],
                "items": [],
            }
        ],
    }


def _write_docx(path: Path, document: dict) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RegulationCreationError("Для создания DOCX требуется python-docx", status_code=500) from exc
    doc = Document()
    title = str(document.get("title") or "Регламент")
    doc.add_heading(title, level=1)
    for section, heading, _section_path, level in _iter_document_sections(
        document.get("sections") or [],
        parent_path=[title],
    ):
        if heading:
            doc.add_heading(heading, level=min(max(level, 2), 9))
        for paragraph in section.get("paragraphs") or []:
            text = str(paragraph or "").strip()
            if text:
                doc.add_paragraph(text)
        for item in section.get("items") or []:
            _add_docx_item(doc, item, level=0)
    doc.save(str(path))


def _iter_document_sections(
    sections: object,
    *,
    parent_path: list[str],
    level: int = 2,
) -> Iterator[tuple[dict, str, list[str], int]]:
    if not isinstance(sections, list):
        return
    for section in sections:
        if not isinstance(section, dict):
            continue
        heading = _document_section_heading(section)
        section_path = [*parent_path, heading] if heading else parent_path
        yield section, heading, section_path, level
        for key in ("sections", "subsections", "children"):
            yield from _iter_document_sections(
                section.get(key),
                parent_path=section_path,
                level=level + 1,
            )


def _document_section_heading(section: dict) -> str:
    heading = str(section.get("title") or "").strip()
    number = str(section.get("number") or "").strip()
    return f"{number} {heading}".strip() if heading else ""


def _document_item_text(item: object) -> str:
    if isinstance(item, dict):
        return str(item.get("text") or item.get("title") or "").strip()
    return str(item or "").strip()


def _add_docx_item(doc: object, item: object, *, level: int) -> None:
    text = _document_item_text(item)
    if text:
        style = "List Bullet" if level <= 0 else "List Bullet 2"
        doc.add_paragraph(text, style=style)
    if isinstance(item, dict):
        children = item.get("items") or item.get("children") or []
        if isinstance(children, list):
            for child in children:
                _add_docx_item(doc, child, level=level + 1)


def _load_creation_attachments(files: list[tuple[str, bytes]]) -> list[dict]:
    loaded: list[dict] = []
    total_chars = 0
    for name, raw in files:
        suffix = Path(name or "").suffix.lower()
        if suffix == ".doc":
            raise RegulationCreationError("Формат DOC не поддерживается. Сохраните файл как DOCX.")
        if suffix not in _CREATION_ATTACH_SUFFIXES:
            raise RegulationCreationError(
                f"Формат «{suffix or 'без расширения'}» не поддерживается. "
                "Допустимо: docx, pdf, md, txt."
            )
        try:
            item = _load_creation_attachment(name, raw)
        except DocumentError as exc:
            raise RegulationCreationError(str(exc)) from exc
        text = str(item.get("text") or "")
        remain = max(0, _MAX_ATTACH_CHARS - total_chars)
        if len(text) > remain:
            item["text"] = text[:remain] + "\n...[текст файла обрезан]"
            text = str(item["text"])
        total_chars += len(text)
        loaded.append(item)
        if total_chars >= _MAX_ATTACH_CHARS:
            break
    return loaded


def _load_creation_attachment(name: str, raw: bytes) -> dict:
    suffix = Path(name or "").suffix.lower()
    if suffix != ".pdf":
        return load_attachment_bytes(name, raw)
    try:
        return load_attachment_bytes(name, raw)
    except DocumentError as exc:
        if "Документ пуст" not in str(exc):
            raise
    with tempfile.TemporaryDirectory(prefix="reg-create-ocr-") as tmp:
        path = Path(tmp) / (Path(name).name or "scan.pdf")
        path.write_bytes(raw)
        is_scan, _page_count = is_scan_pdf(path)
        if not is_scan:
            raise DocumentError("Документ пуст или не удалось извлечь текст.")
        try:
            extracted = extract_pdf_scan(path, work_dir=Path(tmp))
        except RuntimeError as exc:
            raise DocumentError(str(exc)) from exc
    text = compose_regulation_text(
        RegulationParseResult(
            regulationId="reg-create-attachment",
            fileName=Path(name).name or "scan.pdf",
            pageCount=extracted.page_count,
            isScan=extracted.is_scan,
            fragments=[
                {
                    "fragmentId": block.block_id or f"ocr-{index}",
                    "page": block.page,
                    "section": block.section or "",
                    "kind": block.kind,
                    "blockType": block.block_type,
                    "text": block.text,
                    "ocrConfidence": block.confidence,
                }
                for index, block in enumerate(extracted.blocks, start=1)
                if (block.text or "").strip()
            ],
        )
    )
    if not text.strip():
        raise DocumentError("Документ пуст или не удалось извлечь текст.")
    return {
        "name": Path(name).name or "scan.pdf",
        "text": text,
        "kind": "text",
        "mime_type": "application/pdf",
        "data_b64": "",
    }


def _display_user_message(message: str, attachments: list[dict]) -> str:
    names = [str(item.get("name") or "file") for item in attachments]
    if not names:
        return message
    note = "📎 " + ", ".join(names)
    return f"{message}\n\n{note}".strip() if message else note


def _attachments_structured(attachments: list[dict]) -> dict:
    if not attachments:
        return {}
    return {
        "attachments": [
            {
                "name": str(item.get("name") or "file"),
                "shortName": _short_attachment_name(str(item.get("name") or "file")),
            }
            for item in attachments
        ]
    }


def _short_attachment_name(name: str, keep: int = 6) -> str:
    path = Path(name)
    stem = path.stem.replace(" ", "_")
    suffix = path.suffix.lower()
    if len(stem) <= keep:
        return f"{stem}{suffix}"
    return f"{stem[:keep]}...{suffix}"


def _first_interview_object(raw: str) -> dict | None:
    decoder = json.JSONDecoder()
    text = raw or ""
    index = 0
    while index < len(text):
        start = text.find("{", index)
        if start < 0:
            break
        try:
            obj, end = decoder.raw_decode(text, start)
        except json.JSONDecodeError:
            index = start + 1
            continue
        if isinstance(obj, dict) and str(obj.get("status") or "") in {"need_more", "ready"}:
            return obj
        index = end
    return None


def _parse_agent_response(raw: str) -> dict:
    text = raw.strip()
    text = re.sub(r"^```(?:json)?", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    first = _first_interview_object(text)
    if first is not None:
        return first
    try:
        data = json.loads(text)
        return data if isinstance(data, dict) else {"status": "need_more", "message": raw}
    except json.JSONDecodeError:
        try:
            data = ast.literal_eval(text)
            return data if isinstance(data, dict) else {"status": "need_more", "message": raw}
        except (SyntaxError, ValueError):
            pass
        match = re.search(r"\{.*\}", text, flags=re.S)
        if match:
            try:
                data = json.loads(match.group(0))
                return data if isinstance(data, dict) else {"status": "need_more", "message": raw}
            except json.JSONDecodeError:
                try:
                    data = ast.literal_eval(match.group(0))
                    return data if isinstance(data, dict) else {"status": "need_more", "message": raw}
                except (SyntaxError, ValueError):
                    pass
    return {"status": "need_more", "message": raw}


def _quick_answers(value: object) -> list[str]:
    if isinstance(value, list):
        answers = [str(item).strip() for item in value if str(item).strip()]
        answers = [
            item
            for item in answers
            if item.lower() not in {"оставить", "переделать", "оставить это"}
        ]
        if answers:
            return answers[:6]
    return [
        "Опишу действие вручную",
        "Приложу файл с деталями",
        "Это выполняется в Outlook",
        "Это выполняется в 1C",
        "Это выполняется в Excel",
    ]


def _is_force_create_message(message: str) -> bool:
    text = message.strip().lower()
    return "принудительно" in text or "создай регламент" in text and "не хватает" in text


def _session(db: Session, draft: RegulationCreationDraft) -> RegulationCreationSession:
    result = None
    if draft.result_regulation_id:
        from app.services.regulation.storage import get_document

        doc = get_document(db, regulation_id=draft.result_regulation_id, user_id=draft.user_id)
        if doc is not None:
            try:
                result = RegulationParseResult.model_validate(doc.result_json)
            except Exception:
                result = None
        if result is None and isinstance(draft.draft_document_json, dict) and draft.draft_document_json:
            result = _result_from_created_document(
                regulation_id=draft.result_regulation_id,
                filename=Path(draft.result_document_path or "regulation.docx").name,
                document=draft.draft_document_json,
            )
    return RegulationCreationSession(
        draftId=draft.id,
        status=draft.status,
        cursorAgentId=draft.cursor_agent_id,
        latestRunId=draft.latest_run_id,
        positions=[str(item) for item in draft.positions_json or []],
        messages=[
            CreationMessageSchema(
                messageId=item.id,
                draftId=item.draft_id,
                role=item.role,
                content=item.content,
                structured=item.structured_json or {},
                createdAt=item.created_at,
            )
            for item in _messages_for_draft(db, draft.id)
        ],
        resultRegulation=result,
        resultDocument=draft.draft_document_json or {},
        resultDocumentPath=draft.result_document_path,
        sdkAgentId=interview_sdk_agent_id(draft.interview_json),
        createdAt=draft.created_at,
        updatedAt=draft.updated_at,
    )


def _messages_for_draft(db: Session, draft_id: str) -> list[RegulationCreationMessage]:
    return (
        db.query(RegulationCreationMessage)
        .filter(RegulationCreationMessage.draft_id == draft_id)
        .order_by(RegulationCreationMessage.created_at.asc())
        .all()
    )


def _add_message(
    db: Session,
    *,
    draft: RegulationCreationDraft,
    role: str,
    content: str,
    structured: dict | None = None,
) -> None:
    db.add(
        RegulationCreationMessage(
            id=f"reg-create-msg-{uuid4().hex[:12]}",
            draft_id=draft.id,
            user_id=draft.user_id,
            role=role,
            content=_clip_message(content),
            structured_json=structured or {},
        )
    )


def _clip_message(content: str, limit: int = _MESSAGE_CONTENT_LIMIT) -> str:
    text = content or ""
    if len(text) <= limit:
        return text
    return text[: limit - 16] + "\n...[truncated]"


def _get_draft(db: Session, *, user_id: str, draft_id: str) -> RegulationCreationDraft:
    draft = (
        db.query(RegulationCreationDraft)
        .filter(RegulationCreationDraft.id == draft_id, RegulationCreationDraft.user_id == user_id)
        .first()
    )
    if draft is None:
        raise RegulationCreationError("Черновик создания регламента не найден", status_code=404)
    return draft


def _safe_filename(value: str) -> str:
    safe = re.sub(r"[^A-Za-zА-Яа-яЁё0-9._ -]+", " ", value).strip()
    return (safe or "created-regulation")[:120]
