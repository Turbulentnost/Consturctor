from __future__ import annotations

from pathlib import Path

import pytest

from app.services.regulation.docx_extract import extract_docx


def test_extract_docx_keeps_style_runs_and_indent(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    heading = doc.add_paragraph()
    heading_run = heading.add_run("Styled heading")
    heading_run.bold = True
    heading_run.font.size = Pt(16)
    heading_run.font.name = "Calibri"
    heading.style = doc.styles["Heading 1"]

    body = doc.add_paragraph()
    body.paragraph_format.left_indent = Inches(0.5)
    mixed = body.add_run("Regular text and ")
    mixed.font.size = Pt(12)
    mixed.font.name = "Times New Roman"
    bold = body.add_run("bold part")
    bold.bold = True
    bold.font.size = Pt(12)
    bold.font.color.rgb = RGBColor(0x0A, 0x48, 0x3D)

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Value A"
    table.cell(1, 1).text = "Value B"

    after_table = doc.add_paragraph()
    after_table.add_run("After table")

    path = tmp_path / "styled.docx"
    doc.save(path)

    extracted = extract_docx(path)
    texts = [block.text for block in extracted.blocks]

    heading_block = next(block for block in extracted.blocks if "Styled heading" in block.text)
    assert heading_block.block_type == "heading"
    assert heading_block.style_runs
    assert heading_block.style_runs[0]["isBold"]
    assert heading_block.style_runs[0]["fontSize"] == 16
    assert heading_block.font_size == 16

    body_block = next(block for block in extracted.blocks if "Regular text" in block.text)
    assert body_block.location["indentPt"] >= 30
    assert len(body_block.style_runs) >= 2
    assert body_block.style_runs[1]["isBold"]
    assert body_block.style_runs[1]["text"] == "bold part"

    table_index = next(index for index, block in enumerate(extracted.blocks) if block.block_type == "table")
    assert table_index > 0
    assert table_index < len(extracted.blocks) - 1
    assert "After table" in texts
    assert texts.index("After table") > table_index
