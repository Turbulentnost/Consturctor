"""Экспорт текстового отчёта в файл (Word .docx с fallback в Markdown).

Инструмент закрывает разрыв: планы часто требуют "отчёт Word / документ",
а report.* возвращают только текст без файла. Здесь агент отдаёт заголовок,
разделы и (необязательно) таблицу, а инструмент кладёт готовый файл в
песочницу агента и возвращает путь в поле ``file``.

Исполняется локально (без COM). Если python-docx недоступен (например, в
32-bit сборке), формируется .md с тем же содержимым - шаг всё равно получает
``file`` и не застревает.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from app.tools.ac.agent_workspace import AgentWorkspaceResolver, WorkspaceError
from app.tools.ac.base import BaseTool
from app.tools.ac.registry import ToolRegistry
from app.tools.ac.tooling import (
    ToolCallResult,
    ToolDefinition,
    ToolExecutionMode,
    ToolSideEffectLevel,
)


def _as_sections(value: Any) -> list[dict[str, str]]:
    sections: list[dict[str, str]] = []
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                heading = str(item.get("heading") or item.get("title") or "").strip()
                body = str(item.get("body") or item.get("text") or "").strip()
                if heading or body:
                    sections.append({"heading": heading, "body": body})
            elif isinstance(item, str) and item.strip():
                sections.append({"heading": "", "body": item.strip()})
    elif isinstance(value, str) and value.strip():
        sections.append({"heading": "", "body": value.strip()})
    return sections


def _as_table(value: Any) -> tuple[list[str], list[list[str]]]:
    if not isinstance(value, dict):
        return [], []
    headers = [str(h) for h in (value.get("headers") or []) if str(h).strip()]
    rows: list[list[str]] = []
    for row in value.get("rows") or []:
        if isinstance(row, (list, tuple)):
            rows.append([_cell(c) for c in row])
        elif isinstance(row, dict) and headers:
            rows.append([_cell(row.get(h, "")) for h in headers])
    return headers, rows


def _cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


class ReportExportDocumentTool(BaseTool):
    """Собрать отчёт из готовых данных и сохранить файлом (Word/Markdown)."""

    def __init__(self, resolver: AgentWorkspaceResolver) -> None:
        super().__init__(
            ToolDefinition(
                name="report.export_document",
                title="Экспорт отчёта в файл",
                description=(
                    "Сохраняет готовый отчёт файлом в папке агента: Word (.docx), "
                    "если доступен, иначе Markdown (.md). На вход - title, sections "
                    "(heading/body), summary и необязательная table (headers/rows). "
                    "Возвращает file (путь). Используй для 'отчёт Word/документ', "
                    "когда нужен именно файл, а не текст. Текст пиши в sections сам - "
                    "инструмент ничего не выдумывает."
                ),
                side_effect_level=ToolSideEffectLevel.CREATE_DRAFT,
                execution_mode=ToolExecutionMode.LOCAL,
                requires_human_approval=False,
                input_schema={
                    "type": "object",
                    "properties": {
                        "filename": {
                            "type": "string",
                            "description": "Имя файла без пути. Расширение подставится само.",
                        },
                        "title": {"type": "string", "description": "Заголовок документа"},
                        "summary": {"type": "string", "description": "Короткое резюме в начале"},
                        "sections": {
                            "type": "array",
                            "description": "Разделы отчёта: [{heading, body}]. body - готовый текст.",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "heading": {"type": "string"},
                                    "body": {"type": "string"},
                                },
                            },
                        },
                        "table": {
                            "type": "object",
                            "description": "Необязательная таблица: {headers: [...], rows: [[...]]}",
                            "properties": {
                                "headers": {"type": "array", "items": {"type": "string"}},
                                "rows": {"type": "array"},
                            },
                        },
                        "format": {
                            "type": "string",
                            "enum": ["docx", "md"],
                            "description": "Желаемый формат. По умолчанию docx с fallback в md.",
                        },
                    },
                    "required": ["filename"],
                },
                output_schema={"type": "object"},
            )
        )
        self._resolver = resolver

    def execute(self, input_data: dict) -> ToolCallResult:
        raw_name = str(input_data.get("filename") or "").strip() or "report"
        stem = Path(raw_name).stem or "report"
        title = str(input_data.get("title") or stem).strip()
        summary = str(input_data.get("summary") or "").strip()
        sections = _as_sections(input_data.get("sections"))
        headers, rows = _as_table(input_data.get("table"))
        want = str(input_data.get("format") or "docx").strip().lower()

        try:
            workspace = self._resolver.for_agent(
                self._resolver.agent_id_from_input(input_data)
            )
        except WorkspaceError as exc:
            return self._fail("WORKSPACE_ERROR", str(exc))

        if not sections and not summary and not rows:
            return self._fail(
                "REPORT_EMPTY",
                "Нет данных для отчёта: заполни summary/sections по собранным данным.",
            )

        use_docx = want != "md" and _docx_available()
        try:
            if use_docx:
                path = self._write_docx(workspace, stem, title, summary, sections, headers, rows)
                fmt = "docx"
            else:
                path = self._write_md(workspace, stem, title, summary, sections, headers, rows)
                fmt = "md"
        except WorkspaceError as exc:
            return self._fail("WORKSPACE_ERROR", str(exc))
        except Exception as exc:  # noqa: BLE001
            return self._fail("REPORT_WRITE_ERROR", str(exc))

        return ToolCallResult(
            ok=True,
            tool_name=self.definition.name,
            output_data={
                "file": str(path),
                "filename": path.name,
                "path": str(path),
                "format": fmt,
                "section_count": len(sections),
                "generated_at": datetime.now().isoformat(timespec="seconds"),
            },
        )

    def _write_docx(
        self,
        workspace: Any,
        stem: str,
        title: str,
        summary: str,
        sections: list[dict[str, str]],
        headers: list[str],
        rows: list[list[str]],
    ) -> Path:
        from docx import Document

        path = workspace.resolve(f"{stem}.docx")
        document = Document()
        document.add_heading(title, level=0)
        if summary:
            document.add_paragraph(summary)
        for section in sections:
            if section["heading"]:
                document.add_heading(section["heading"], level=1)
            if section["body"]:
                for line in section["body"].splitlines() or [section["body"]]:
                    document.add_paragraph(line)
        if headers and rows:
            table = document.add_table(rows=1, cols=len(headers))
            for idx, head in enumerate(headers):
                table.rows[0].cells[idx].text = head
            for row in rows:
                cells = table.add_row().cells
                for idx in range(len(headers)):
                    cells[idx].text = row[idx] if idx < len(row) else ""
        document.save(str(path))
        return path

    def _write_md(
        self,
        workspace: Any,
        stem: str,
        title: str,
        summary: str,
        sections: list[dict[str, str]],
        headers: list[str],
        rows: list[list[str]],
    ) -> Path:
        path = workspace.resolve(f"{stem}.md")
        lines: list[str] = [f"# {title}", ""]
        if summary:
            lines += [summary, ""]
        for section in sections:
            if section["heading"]:
                lines += [f"## {section['heading']}", ""]
            if section["body"]:
                lines += [section["body"], ""]
        if headers and rows:
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows:
                cells = [row[i] if i < len(row) else "" for i in range(len(headers))]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def _fail(self, error_type: str, message: str) -> ToolCallResult:
        return ToolCallResult(
            ok=False,
            tool_name=self.definition.name,
            error_type=error_type,
            error_message=message,
        )


def _docx_available() -> bool:
    import importlib.util

    return importlib.util.find_spec("docx") is not None


def register_document_tools(
    registry: ToolRegistry,
    resolver: AgentWorkspaceResolver,
    *,
    skip_existing: bool = False,
) -> None:
    """Зарегистрировать инструмент экспорта отчёта в файл."""
    tool = ReportExportDocumentTool(resolver)
    if skip_existing and registry.has_tool(tool.definition.name):
        return
    registry.register(tool)
